"""Unit tests for export_handler.py — FE-UI-028 acceptance criteria."""

from __future__ import annotations

import json
from datetime import datetime, timezone
from unittest.mock import MagicMock, patch

import pytest

from careervp.handlers.export_handler import (
    _build_docx,
    _write_and_presign,
    lambda_handler,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

USER_ID = 'user-sub-123'
JOB_ID = 'job-abc'


def _make_event(
    job_id: str = JOB_ID,
    module_type: str = 'vpr',
    fmt: str = 'docx',
    method: str = 'GET',
    user_id: str = USER_ID,
    include_auth: bool = True,
) -> dict:
    event: dict = {
        'httpMethod': method,
        'path': f'/jobs/{job_id}/artifacts/{module_type}/export',
        'pathParameters': {'jobId': job_id, 'moduleType': module_type},
        'queryStringParameters': {'format': fmt},
        'headers': {},
        'body': None,
    }
    if include_auth:
        event['requestContext'] = {'authorizer': {'claims': {'sub': user_id}}}
    return event


def _mock_vpr_s3(monkeypatch: pytest.MonkeyPatch, data: dict | None = None) -> MagicMock:
    vpr_data = data or {'summary': 'Great candidate', 'skills': 'Python, AWS'}
    mock_s3 = MagicMock()
    mock_s3.get_object.return_value = {'Body': MagicMock(read=lambda: json.dumps(vpr_data).encode())}
    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')
    return mock_s3


def _mock_dynamo_item(item: dict | None) -> MagicMock:
    table = MagicMock()
    table.get_item.return_value = {'Item': item} if item is not None else {}
    resource = MagicMock()
    resource.Table.return_value = table
    return resource


# ---------------------------------------------------------------------------
# AC-005: ?format=pdf → 501
# ---------------------------------------------------------------------------


def test_pdf_format_returns_501(monkeypatch: pytest.MonkeyPatch) -> None:
    event = _make_event(fmt='pdf')
    response = lambda_handler(event, None)
    assert response['statusCode'] == 501
    body = json.loads(response['body'])
    assert 'pdf' in body['message'].lower() or 'not yet' in body['message'].lower()


# ---------------------------------------------------------------------------
# AC-006: unknown format → 400
# ---------------------------------------------------------------------------


def test_unknown_format_returns_400(monkeypatch: pytest.MonkeyPatch) -> None:
    event = _make_event(fmt='xlsx')
    response = lambda_handler(event, None)
    assert response['statusCode'] == 400
    body = json.loads(response['body'])
    assert 'unsupported' in body['message'].lower() or 'format' in body['message'].lower()


def test_missing_format_returns_400(monkeypatch: pytest.MonkeyPatch) -> None:
    event = _make_event()
    event['queryStringParameters'] = {}
    response = lambda_handler(event, None)
    assert response['statusCode'] == 400


# ---------------------------------------------------------------------------
# AC-007: unknown moduleType → 400
# ---------------------------------------------------------------------------


def test_unknown_module_type_returns_400(monkeypatch: pytest.MonkeyPatch) -> None:
    event = _make_event(module_type='unknown_module', fmt='docx')
    response = lambda_handler(event, None)
    assert response['statusCode'] == 400
    body = json.loads(response['body'])
    assert 'unsupported' in body['message'].lower() or 'module' in body['message'].lower()


# ---------------------------------------------------------------------------
# OPTIONS pre-flight
# ---------------------------------------------------------------------------


def test_options_returns_200() -> None:
    event = _make_event(method='OPTIONS')
    response = lambda_handler(event, None)
    assert response['statusCode'] == 200


# ---------------------------------------------------------------------------
# AC-008: vpr artifact missing → 404
# ---------------------------------------------------------------------------


def test_vpr_missing_artifact_returns_404(monkeypatch: pytest.MonkeyPatch) -> None:
    import botocore.exceptions

    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    mock_s3 = MagicMock()
    error_response = {'Error': {'Code': 'NoSuchKey', 'Message': 'Not found'}}
    mock_s3.get_object.side_effect = botocore.exceptions.ClientError(error_response, 'GetObject')

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='vpr'), None)

    assert response['statusCode'] == 404
    body = json.loads(response['body'])
    assert 'not found' in body['message'].lower()


# ---------------------------------------------------------------------------
# AC-009: cover_letter artifact missing → 404
# ---------------------------------------------------------------------------


def test_cover_letter_missing_artifact_returns_404(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('ARTIFACTS_TABLE_NAME', 'artifacts-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    mock_resource = _mock_dynamo_item(None)

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        response = lambda_handler(_make_event(module_type='cover_letter'), None)

    assert response['statusCode'] == 404


# ---------------------------------------------------------------------------
# AC-010: interview_prep artifact missing → 404
# ---------------------------------------------------------------------------


def test_interview_prep_missing_artifact_returns_404(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('ARTIFACTS_TABLE_NAME', 'artifacts-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    mock_resource = _mock_dynamo_item(None)

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        response = lambda_handler(_make_event(module_type='interview_prep'), None)

    assert response['statusCode'] == 404


# ---------------------------------------------------------------------------
# AC-011: cv_tailored artifact missing → 404
# ---------------------------------------------------------------------------


def test_cv_tailored_missing_artifact_returns_404(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('TABLE_NAME', 'main-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    mock_resource = _mock_dynamo_item(None)

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        response = lambda_handler(_make_event(module_type='cv_tailored'), None)

    assert response['statusCode'] == 404


# ---------------------------------------------------------------------------
# AC-001: vpr → 200 with download_url + expires_at
# ---------------------------------------------------------------------------


def test_vpr_export_returns_200_with_download_url(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    vpr_data = {'summary': 'Strong candidate', 'skills': 'Python, AWS'}
    mock_s3 = MagicMock()
    mock_s3.get_object.return_value = {'Body': MagicMock(read=lambda: json.dumps(vpr_data).encode())}
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/presigned'

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='vpr'), None)

    assert response['statusCode'] == 200
    body = json.loads(response['body'])
    assert 'download_url' in body
    assert 'expires_at' in body
    assert body['download_url'] == 'https://s3.example.com/presigned'

    # AC-013: presigned URL key must target exports/vpr/{jobId}/{jobId}.docx
    put_call_args = mock_s3.put_object.call_args
    assert f'exports/vpr/{JOB_ID}/{JOB_ID}.docx' in str(put_call_args)


# ---------------------------------------------------------------------------
# AC-002: cover_letter → 200
# ---------------------------------------------------------------------------


def test_cover_letter_export_returns_200(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('ARTIFACTS_TABLE_NAME', 'artifacts-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    item = {'applicationId': USER_ID, 'artifactId': JOB_ID, 'cover_letter': {'full_text': 'Dear Hiring Manager...'}}
    mock_resource = _mock_dynamo_item(item)

    mock_s3 = MagicMock()
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/cover-letter'

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='cover_letter'), None)

    assert response['statusCode'] == 200
    body = json.loads(response['body'])
    assert 'download_url' in body
    assert 'expires_at' in body


# ---------------------------------------------------------------------------
# AC-003: interview_prep → 200 with Q&A pairs in docx
# ---------------------------------------------------------------------------


def test_interview_prep_export_returns_200(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('ARTIFACTS_TABLE_NAME', 'artifacts-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    item = {
        'applicationId': USER_ID,
        'artifactId': f'ARTIFACT#INTERVIEW_PREP#{JOB_ID}',
        'interview_prep': {
            'questions': [
                {'question': 'Tell me about yourself', 'answer': 'I am a developer'},
                {'question': 'Strengths?', 'answer': 'Fast learner'},
            ]
        },
    }
    mock_resource = _mock_dynamo_item(item)

    mock_s3 = MagicMock()
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/interview-prep'

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='interview_prep'), None)

    assert response['statusCode'] == 200


# ---------------------------------------------------------------------------
# AC-004: cv_tailored → 200 with cv_sections + tailored_cv
# ---------------------------------------------------------------------------


def test_cv_tailored_export_returns_200(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('TABLE_NAME', 'main-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    item = {
        'pk': USER_ID,
        'sk': f'ARTIFACT#CV_TAILORED#{JOB_ID}',
        'cv_sections': {'experience': 'Led team of 5', 'skills': 'Python'},
        'tailored_cv': 'Experienced engineer targeting ML roles.',
    }
    mock_resource = _mock_dynamo_item(item)

    mock_s3 = MagicMock()
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/cv-tailored'

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='cv_tailored'), None)

    assert response['statusCode'] == 200
    body = json.loads(response['body'])
    assert 'download_url' in body


# ---------------------------------------------------------------------------
# AC-013: expires_at is within ~3600 s of now
# ---------------------------------------------------------------------------


def test_expires_at_is_approximately_one_hour_from_now(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    vpr_data = {'summary': 'Test'}
    mock_s3 = MagicMock()
    mock_s3.get_object.return_value = {'Body': MagicMock(read=lambda: json.dumps(vpr_data).encode())}
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/presigned'

    before = datetime.now(timezone.utc)

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='vpr'), None)

    assert response['statusCode'] == 200
    body = json.loads(response['body'])
    expires_at = datetime.strptime(body['expires_at'], '%Y-%m-%dT%H:%M:%SZ').replace(tzinfo=timezone.utc)
    delta = (expires_at - before).total_seconds()
    assert 3599 <= delta <= 3601


# ---------------------------------------------------------------------------
# AC-015: S3 key follows exports/{moduleType}/{jobId}/{jobId}.docx
# ---------------------------------------------------------------------------


def test_s3_key_pattern_is_correct(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    vpr_data = {'section': 'value'}
    mock_s3 = MagicMock()
    mock_s3.get_object.return_value = {'Body': MagicMock(read=lambda: json.dumps(vpr_data).encode())}
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/presigned'

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        lambda_handler(_make_event(module_type='vpr', job_id='my-job'), None)

    put_kwargs = mock_s3.put_object.call_args.kwargs
    assert put_kwargs['Key'] == 'exports/vpr/my-job/my-job.docx'
    assert put_kwargs['Bucket'] == 'artifacts-bucket'


# ---------------------------------------------------------------------------
# _build_docx unit tests (DOCX content structure)
# ---------------------------------------------------------------------------


def test_build_docx_vpr_adds_sections() -> None:
    data = {'summary': 'Top candidate', 'technical_skills': 'Python, Go'}
    doc = _build_docx('vpr', data)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Top candidate' in full_text


def test_build_docx_cover_letter_uses_full_text() -> None:
    data = {'cover_letter': {'full_text': 'Dear Hiring Manager, I am excited...'}}
    doc = _build_docx('cover_letter', data)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Dear Hiring Manager' in full_text


def test_build_docx_cover_letter_falls_back_to_text_field() -> None:
    data = {'cover_letter': {'text': 'Fallback text content'}}
    doc = _build_docx('cover_letter', data)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Fallback text content' in full_text


def test_build_docx_interview_prep_renders_qa_pairs() -> None:
    data = {
        'interview_prep': {
            'questions': [
                {'question': 'Why this role?', 'answer': 'Passion for AI'},
                {'question': 'Biggest strength?', 'answer': 'Problem solving'},
            ]
        }
    }
    doc = _build_docx('interview_prep', data)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Why this role?' in full_text
    assert 'Passion for AI' in full_text
    assert '1.' in full_text
    assert '2.' in full_text


def test_build_docx_cv_tailored_includes_sections_and_summary() -> None:
    data = {
        'cv_sections': {'experience': 'Led ML projects'},
        'tailored_cv': 'Strong ML background targeting senior roles.',
    }
    doc = _build_docx('cv_tailored', data)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Led ML projects' in full_text
    assert 'Strong ML background' in full_text


# ---------------------------------------------------------------------------
# _write_and_presign produces a valid DOCX buffer (AC-015)
# ---------------------------------------------------------------------------


def test_write_and_presign_puts_valid_docx_to_s3() -> None:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph('Test content')

    mock_s3 = MagicMock()
    mock_s3.put_object.return_value = {}
    mock_s3.generate_presigned_url.return_value = 'https://s3.example.com/test'

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        url, expires_at = _write_and_presign('my-bucket', 'exports/vpr/job1/job1.docx', doc)

    assert url == 'https://s3.example.com/test'
    put_kwargs = mock_s3.put_object.call_args.kwargs
    body_bytes: bytes = put_kwargs['Body']
    # DOCX is a ZIP — must start with PK magic bytes
    assert body_bytes[:2] == b'PK', 'Uploaded bytes are not a valid DOCX/ZIP'


# ---------------------------------------------------------------------------
# snake_case path params backward-compat (existing test coverage)
# ---------------------------------------------------------------------------


def test_snake_case_path_params_are_supported(monkeypatch: pytest.MonkeyPatch) -> None:
    """Handler must also accept job_id / module_type for backward compatibility."""
    event = _make_event()
    event['pathParameters'] = {'job_id': JOB_ID, 'module_type': 'unknown_module'}
    response = lambda_handler(event, None)
    # unknown_module → 400 (not a crash)
    assert response['statusCode'] == 400


# ---------------------------------------------------------------------------
# OPTIONS CORS header behaviour
# ---------------------------------------------------------------------------


def test_options_allowed_origin_returns_cors_header(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr('careervp.handlers.cors_utils._ALLOWED_ORIGINS', {'https://allowed.example.com'})
    event = _make_event(method='OPTIONS')
    event['headers'] = {'origin': 'https://allowed.example.com'}
    response = lambda_handler(event, None)
    assert response['statusCode'] == 200
    assert response['headers'].get('Access-Control-Allow-Origin') == 'https://allowed.example.com'


def test_options_disallowed_origin_has_no_cors_header(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr('careervp.handlers.cors_utils._ALLOWED_ORIGINS', {'https://allowed.example.com'})
    event = _make_event(method='OPTIONS')
    event['headers'] = {'origin': 'https://malicious.example.com'}
    response = lambda_handler(event, None)
    assert response['statusCode'] == 200
    assert 'Access-Control-Allow-Origin' not in response['headers']


# ---------------------------------------------------------------------------
# Data source read failure → 500
# ---------------------------------------------------------------------------


def test_vpr_data_source_non_404_error_returns_500(monkeypatch: pytest.MonkeyPatch) -> None:
    import botocore.exceptions

    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    mock_s3 = MagicMock()
    mock_s3.get_object.side_effect = botocore.exceptions.ClientError({'Error': {'Code': '403', 'Message': 'AccessDenied'}}, 'GetObject')

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='vpr'), None)

    assert response['statusCode'] == 500


@pytest.mark.parametrize('module_type', ['cover_letter', 'interview_prep', 'cv_tailored'])
def test_dynamo_read_failure_returns_500(monkeypatch: pytest.MonkeyPatch, module_type: str) -> None:
    import botocore.exceptions

    monkeypatch.setenv('ARTIFACTS_TABLE_NAME', 'artifacts-table')
    monkeypatch.setenv('TABLE_NAME', 'main-table')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    mock_table = MagicMock()
    mock_table.get_item.side_effect = botocore.exceptions.ClientError(
        {'Error': {'Code': '400', 'Message': 'ProvisionedThroughputExceededException'}}, 'GetItem'
    )
    mock_resource = MagicMock()
    mock_resource.Table.return_value = mock_table

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.resource.return_value = mock_resource
        response = lambda_handler(_make_event(module_type=module_type), None)

    assert response['statusCode'] == 500


# ---------------------------------------------------------------------------
# S3 upload failure → 500
# ---------------------------------------------------------------------------


def test_s3_upload_failure_returns_500(monkeypatch: pytest.MonkeyPatch) -> None:
    import botocore.exceptions

    monkeypatch.setenv('VPR_RESULTS_BUCKET_NAME', 'vpr-bucket')
    monkeypatch.setenv('ARTIFACTS_BUCKET_NAME', 'artifacts-bucket')

    vpr_data = {'summary': 'Test candidate'}
    mock_s3 = MagicMock()
    mock_s3.get_object.return_value = {'Body': MagicMock(read=lambda: json.dumps(vpr_data).encode())}
    mock_s3.put_object.side_effect = botocore.exceptions.ClientError({'Error': {'Code': '503', 'Message': 'SlowDown'}}, 'PutObject')

    with patch('careervp.handlers.export_handler.boto3') as mock_boto3:
        mock_boto3.client.return_value = mock_s3
        response = lambda_handler(_make_event(module_type='vpr'), None)

    assert response['statusCode'] == 500
