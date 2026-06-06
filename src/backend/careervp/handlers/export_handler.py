"""
Lambda handler for the Export API endpoint.
Generates DOCX files from stored artifacts and returns presigned S3 download URLs.
Endpoint: GET /jobs/{jobId}/artifacts/{moduleType}/export?format=docx|pdf
"""

from __future__ import annotations

import io
import json
import os
from datetime import datetime, timedelta, timezone
from http import HTTPStatus
from typing import Any

import boto3
import botocore.exceptions
from docx import Document
from docx.document import Document as DocxDocument

from careervp.handlers.cors_utils import get_cors_headers, set_request_origin
from careervp.handlers.utils.observability import logger

INTERVIEW_PREP_SORT_KEY_PREFIX = 'ARTIFACT#INTERVIEW_PREP#'
VALID_MODULE_TYPES = frozenset({'vpr', 'cover_letter', 'interview_prep', 'cv_tailored'})
PRESIGNED_URL_TTL = 3600


class ArtifactNotFoundError(Exception):
    pass


def lambda_handler(event: dict[str, Any], context: Any) -> dict[str, Any]:
    """Route export requests."""
    _ = context
    set_request_origin(event)
    method = str(event.get('httpMethod', 'GET')).upper()

    if method == 'OPTIONS':
        return _json_response(HTTPStatus.OK, {'status': 'ok'})

    if method == 'GET':
        return _handle_export(event)

    return _json_response(HTTPStatus.METHOD_NOT_ALLOWED, {'error': 'Method not allowed'})


def _handle_export(event: dict[str, Any]) -> dict[str, Any]:
    """Handle GET /jobs/{jobId}/artifacts/{moduleType}/export."""
    path_params = event.get('pathParameters') or {}
    job_id = str(path_params.get('jobId') or path_params.get('job_id') or '').strip()
    module_type = str(path_params.get('moduleType') or path_params.get('module_type') or '').strip()

    query_params = event.get('queryStringParameters') or {}
    export_format = str(query_params.get('format', '')).strip().lower()

    try:
        user_id = str(event['requestContext']['authorizer']['claims']['sub']).strip()
    except (KeyError, TypeError):
        return _json_response(HTTPStatus.UNAUTHORIZED, {'error': 'Unauthorized'})

    if not user_id:
        return _json_response(HTTPStatus.UNAUTHORIZED, {'error': 'Unauthorized'})

    if export_format == 'pdf':
        return _json_response(
            HTTPStatus.NOT_IMPLEMENTED,
            {'message': 'PDF export is not yet available.'},
        )
    if export_format != 'docx':
        return _json_response(
            HTTPStatus.BAD_REQUEST,
            {'message': 'Unsupported format. Use docx.'},
        )

    if module_type not in VALID_MODULE_TYPES:
        return _json_response(
            HTTPStatus.BAD_REQUEST,
            {'message': 'Unsupported module type.'},
        )

    logger.info('Export requested', job_id=job_id, module_type=module_type, user_id=user_id)

    try:
        data = _read_artifact(module_type, job_id, user_id)
    except ArtifactNotFoundError:
        return _json_response(HTTPStatus.NOT_FOUND, {'message': 'Artifact not found.'})
    except Exception as exc:
        logger.error('Failed to read artifact', exc_info=exc)
        return _json_response(
            HTTPStatus.INTERNAL_SERVER_ERROR,
            {'message': 'Export failed. Please try again.'},
        )

    try:
        doc = _build_docx(module_type, data)
        artifacts_bucket = os.environ['ARTIFACTS_BUCKET_NAME']
        key = f'exports/{module_type}/{job_id}/{job_id}.docx'
        download_url, expires_at = _write_and_presign(artifacts_bucket, key, doc)
    except Exception as exc:
        logger.error('Failed to generate or upload export', exc_info=exc)
        return _json_response(
            HTTPStatus.INTERNAL_SERVER_ERROR,
            {'message': 'Export failed. Please try again.'},
        )

    return _json_response(HTTPStatus.OK, {'download_url': download_url, 'expires_at': expires_at})


def _read_artifact(module_type: str, job_id: str, user_id: str) -> Any:
    if module_type == 'vpr':
        return _read_vpr(job_id)
    if module_type == 'cover_letter':
        return _read_cover_letter(job_id, user_id)
    if module_type == 'interview_prep':
        return _read_interview_prep(job_id, user_id)
    return _read_cv_tailored(job_id, user_id)


def _read_vpr(job_id: str) -> dict[str, Any]:
    s3 = boto3.client('s3')
    bucket = os.environ['VPR_RESULTS_BUCKET_NAME']
    key = f'results/{job_id}.json'
    try:
        response = s3.get_object(Bucket=bucket, Key=key)
        return dict(json.loads(response['Body'].read()))
    except botocore.exceptions.ClientError as exc:
        if exc.response['Error']['Code'] in ('NoSuchKey', '404'):
            raise ArtifactNotFoundError(f'VPR artifact not found: {job_id}') from exc
        raise


def _read_cover_letter(job_id: str, user_id: str) -> dict[str, Any]:
    dynamodb = boto3.resource('dynamodb')
    table = dynamodb.Table(os.environ['ARTIFACTS_TABLE_NAME'])
    artifact_key = f'ARTIFACT#COVER_LETTER#{job_id}'
    response = table.get_item(Key={'applicationId': user_id, 'artifactId': artifact_key})
    item = response.get('Item')
    if not item:
        raise ArtifactNotFoundError(f'Cover letter artifact not found: {job_id}')
    return dict(item)


def _read_interview_prep(job_id: str, user_id: str) -> dict[str, Any]:
    dynamodb = boto3.resource('dynamodb')
    table = dynamodb.Table(os.environ['ARTIFACTS_TABLE_NAME'])
    artifact_id = f'{INTERVIEW_PREP_SORT_KEY_PREFIX}{job_id}'
    response = table.get_item(Key={'applicationId': user_id, 'artifactId': artifact_id})
    item = response.get('Item')
    if not item:
        raise ArtifactNotFoundError(f'Interview prep artifact not found: {job_id}')
    return dict(item)


def _read_cv_tailored(job_id: str, user_id: str) -> dict[str, Any]:
    dynamodb = boto3.resource('dynamodb')
    table = dynamodb.Table(os.environ['TABLE_NAME'])
    sk = f'ARTIFACT#CV_TAILORED#{job_id}'
    response = table.get_item(Key={'pk': user_id, 'sk': sk})
    item = response.get('Item')
    if not item:
        raise ArtifactNotFoundError(f'CV tailored artifact not found: {job_id}')
    return dict(item)


def _build_docx(module_type: str, data: Any) -> DocxDocument:
    doc = Document()
    builders = {
        'vpr': _fill_vpr,
        'cover_letter': _fill_cover_letter,
        'interview_prep': _fill_interview_prep,
        'cv_tailored': _fill_cv_tailored,
    }
    builders[module_type](doc, data)
    return doc


def _render_value_as_text(value: Any, depth: int = 0) -> str:
    """Recursively convert any value to human-readable text."""
    if value is None:
        return ''
    if isinstance(value, bool):
        return 'Yes' if value else 'No'
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        parts = [_render_value_as_text(item, depth + 1) for item in value]
        return '\n'.join(p for p in parts if p)
    if isinstance(value, dict):
        parts = []
        for k, v in value.items():
            t = _render_value_as_text(v, depth + 1)
            if t:
                label = str(k).replace('_', ' ').replace('-', ' ').title()
                parts.append(f'{label}: {t}')
        return '\n'.join(parts)
    return str(value)


_VPR_SKIP_KEYS = frozenset(
    {
        'applicationId',
        'application_id',
        'pk',
        'sk',
        'status',
        'user_id',
        'userId',
        'created_at',
        'updated_at',
    }
)


def _add_bold_kv(doc: DocxDocument, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f'{str(key).replace("_", " ").title()}: ').bold = True
    p.add_run(value)


def _add_vpr_list(doc: DocxDocument, items: list[Any]) -> None:
    for item in items:
        if isinstance(item, str) and item.strip():
            doc.add_paragraph(f'• {item.strip()}')
        elif isinstance(item, dict):
            for k, v in item.items():
                text = _render_value_as_text(v)
                if text:
                    _add_bold_kv(doc, k, text)
        elif item is not None:
            doc.add_paragraph(f'• {_render_value_as_text(item)}')


def _add_vpr_dict(doc: DocxDocument, mapping: dict[str, Any]) -> None:
    for k, v in mapping.items():
        text = _render_value_as_text(v)
        if text:
            _add_bold_kv(doc, k, text)


def _fill_vpr(doc: DocxDocument, data: Any) -> None:
    doc.add_heading('Value Proposition Report', 0)
    if not isinstance(data, dict):
        return
    for section_key, section_value in data.items():
        if section_key in _VPR_SKIP_KEYS:
            continue
        section_title = str(section_key).replace('_', ' ').replace('-', ' ').title()
        doc.add_heading(section_title, level=1)
        if isinstance(section_value, str) and section_value.strip():
            doc.add_paragraph(section_value.strip())
        elif isinstance(section_value, (int, float, bool)):
            doc.add_paragraph(str(section_value))
        elif isinstance(section_value, list):
            _add_vpr_list(doc, section_value)
        elif isinstance(section_value, dict):
            _add_vpr_dict(doc, section_value)


def _fill_cover_letter(doc: DocxDocument, data: Any) -> None:
    doc.add_heading('Cover Letter', 0)
    cover_letter = data.get('cover_letter') or {} if isinstance(data, dict) else {}
    text = (cover_letter.get('full_text') or cover_letter.get('text') or '') if isinstance(cover_letter, dict) else ''
    doc.add_paragraph(str(text))


def _fill_interview_prep(doc: DocxDocument, data: Any) -> None:
    doc.add_heading('Interview Preparation', 0)
    prep = data.get('interview_prep') or {} if isinstance(data, dict) else {}
    questions = prep.get('questions') or [] if isinstance(prep, dict) else []
    for i, qa in enumerate(questions, start=1):
        if isinstance(qa, dict):
            doc.add_paragraph(f'{i}. Q: {qa.get("question", "")}')
            doc.add_paragraph(f'   A: {qa.get("answer", "")}')
        else:
            doc.add_paragraph(f'{i}. {qa}')


def _cv_add_contact(doc: DocxDocument, cv: dict[str, Any]) -> None:
    contact = cv.get('contact') or {}
    if not isinstance(contact, dict):
        return
    name = contact.get('name') or ''
    if name:
        doc.add_heading(str(name), level=1)
    parts = [contact.get(f) or '' for f in ('email', 'phone', 'linkedin', 'location')]
    line = ' | '.join(p for p in parts if p)
    if line:
        doc.add_paragraph(line)


def _cv_add_experience(doc: DocxDocument, experience: list[Any]) -> None:
    doc.add_heading('Professional Experience', level=2)
    for exp in experience:
        if not isinstance(exp, dict):
            continue
        title = exp.get('title') or ''
        company = exp.get('company') or ''
        start = exp.get('start_date') or ''
        end = 'Present' if exp.get('is_current') else (exp.get('end_date') or '')
        header = f'{title} | {company}' if title and company else (title or company)
        date_range = f'{start} – {end}' if start else end
        p = doc.add_paragraph()
        p.add_run(header).bold = True
        if date_range:
            p.add_run(f'  {date_range}')
        for bullet in exp.get('bullets') or []:
            text = bullet.get('text') or '' if isinstance(bullet, dict) else str(bullet)
            if text:
                doc.add_paragraph(f'• {text}')


def _cv_add_education(doc: DocxDocument, education: list[Any]) -> None:
    doc.add_heading('Education', level=2)
    for edu in education:
        if not isinstance(edu, dict):
            continue
        degree = edu.get('degree') or ''
        field = edu.get('field') or ''
        institution = edu.get('institution') or ''
        grad = edu.get('graduation_date') or ''
        line = f'{degree} in {field} | {institution}' if degree and field else institution
        if grad:
            line += f' | {grad}'
        doc.add_paragraph(line)


def _cv_add_certifications(doc: DocxDocument, certs: list[Any]) -> None:
    doc.add_heading('Certifications', level=2)
    for cert in certs:
        if not isinstance(cert, dict):
            continue
        parts = [cert.get(f) or '' for f in ('name', 'issuer', 'date')]
        line = ' | '.join(p for p in parts if p)
        if line:
            doc.add_paragraph(f'• {line}')


def _cv_add_sections_dict(doc: DocxDocument, cv: dict[str, Any]) -> None:
    _cv_add_contact(doc, cv)
    summary = cv.get('summary') or ''
    if summary:
        doc.add_heading('Professional Summary', level=2)
        doc.add_paragraph(str(summary))
    skills = cv.get('skills') or {}
    if isinstance(skills, dict):
        technical = skills.get('technical') or []
        if technical:
            doc.add_heading('Core Competencies', level=2)
            doc.add_paragraph(' | '.join(str(s) for s in technical))
    experience = cv.get('experience') or []
    if experience:
        _cv_add_experience(doc, experience)
    education = cv.get('education') or []
    if education:
        _cv_add_education(doc, education)
    certs = cv.get('certifications') or []
    if certs:
        _cv_add_certifications(doc, certs)


def _cv_add_sections_list(doc: DocxDocument, sections: list[Any]) -> None:
    for section in sections:
        if not isinstance(section, dict):
            continue
        title = section.get('title') or section.get('name') or ''
        content = section.get('content') or section.get('text') or ''
        if title:
            doc.add_heading(str(title), level=2)
        if content:
            doc.add_paragraph(str(content))


def _fill_cv_tailored(doc: DocxDocument, data: Any) -> None:
    doc.add_heading('Tailored CV', 0)
    if not isinstance(data, dict):
        return

    # Result may be nested under 'result' key (DynamoDB artifact format)
    raw_payload = data.get('result') or data
    payload: dict[str, Any] = raw_payload if isinstance(raw_payload, dict) else data

    cv_sections = payload.get('cv_sections')
    tailored_cv = payload.get('tailored_cv') or ''

    if isinstance(cv_sections, dict) and cv_sections:
        _cv_add_sections_dict(doc, cv_sections)
    elif isinstance(cv_sections, list):
        _cv_add_sections_list(doc, cv_sections)

    if tailored_cv:
        if cv_sections:
            doc.add_heading('Raw Text', level=2)
        doc.add_paragraph(str(tailored_cv))


def _write_and_presign(bucket: str, key: str, doc: DocxDocument) -> tuple[str, str]:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    s3 = boto3.client('s3')
    s3.put_object(
        Bucket=bucket,
        Key=key,
        Body=buffer.getvalue(),
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

    download_url: str = s3.generate_presigned_url(
        'get_object',
        Params={'Bucket': bucket, 'Key': key},
        ExpiresIn=PRESIGNED_URL_TTL,
    )

    expires_dt = datetime.now(timezone.utc) + timedelta(seconds=PRESIGNED_URL_TTL)
    expires_at = expires_dt.strftime('%Y-%m-%dT%H:%M:%SZ')

    return download_url, expires_at


def _json_response(status_code: HTTPStatus, body: dict[str, Any]) -> dict[str, Any]:
    headers = get_cors_headers(None)
    headers['Content-Type'] = 'application/json'
    return {
        'statusCode': status_code.value,
        'headers': headers,
        'body': json.dumps(body),
    }


__all__ = ['lambda_handler']
